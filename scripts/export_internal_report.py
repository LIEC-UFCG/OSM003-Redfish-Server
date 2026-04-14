#!/usr/bin/env python3
"""Export internal Markdown reports into shareable DOCX files.

This script keeps Markdown as the working source of truth and produces a
formatted DOCX file that can be shared with management.

It also generates a compact executive summary file alongside the full report.
"""

from __future__ import annotations

import argparse
import datetime as dt
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_INPUT_DIR = ROOT / "internal_reports"
DEFAULT_OUTPUT_DIR = DEFAULT_INPUT_DIR / "exported"


def parse_args() -> argparse.Namespace:
    """Parse command-line arguments for report export."""
    parser = argparse.ArgumentParser(
        description="Convert an internal Markdown report into a shareable DOCX file."
    )
    parser.add_argument(
        "input",
        nargs="?",
        help="Path to a Markdown report. Defaults to the latest .md file in internal_reports/.",
    )
    parser.add_argument(
        "--output-dir",
        default=str(DEFAULT_OUTPUT_DIR),
        help="Directory where exported files will be written.",
    )
    parser.add_argument(
        "--title",
        default="Relatorio Interno de Evolucao",
        help="Fallback title used when the Markdown file does not define a H1.",
    )
    return parser.parse_args()


def find_latest_report(directory: Path) -> Path:
    """Return the most recently modified Markdown report in a directory."""
    candidates = sorted(
        directory.glob("*.md"),
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )
    if not candidates:
        raise FileNotFoundError(f"No Markdown reports found in {directory}")
    return candidates[0]


def parse_report_metadata(text: str) -> dict[str, str]:
    """Extract simple key/value metadata from the report header."""
    metadata: dict[str, str] = {}
    patterns = {
        "month": r"^Mes/Ano:\s*(.+)$",
        "responsible": r"^Responsavel:\s*(.+)$",
        "team": r"^Equipe/Area:\s*(.+)$",
    }
    for key, pattern in patterns.items():
        match = re.search(pattern, text, flags=re.MULTILINE)
        metadata[key] = match.group(1).strip() if match else ""
    return metadata


def extract_document_title_and_body(text: str, fallback_title: str) -> tuple[str, str]:
    """Extract the first Markdown H1 as title and return body without it."""
    lines = text.splitlines()
    for index, line in enumerate(lines):
        if line.startswith("# "):
            title = line[2:].strip() or fallback_title
            body = "\n".join(lines[:index] + lines[index + 1 :]).lstrip()
            return title, body
    return fallback_title, text


def build_table_of_contents(text: str) -> list[tuple[int, str, str]]:
    """Build a small table of contents from Markdown headings."""
    toc: list[tuple[int, str, str]] = []
    for line in text.splitlines():
        if not line.startswith("## ") and not line.startswith("### "):
            continue
        if line.startswith("## "):
            level = 1
            heading = line[3:].strip()
        else:
            level = 2
            heading = line[4:].strip()
        anchor = re.sub(r"[^a-z0-9]+", "-", heading.lower()).strip("-")
        toc.append((level, heading, anchor))
    return toc


def set_document_defaults(document: Document) -> None:
    """Apply consistent margins and default font settings."""
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    """Set a cell's text while keeping formatting simple and consistent."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)


def add_horizontal_rule(document: Document) -> None:
    """Insert a horizontal line separator into the document."""
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run._r.append(OxmlElement("w:br"))
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    border.append(bottom)
    paragraph._p.get_or_add_pPr().append(border)


def add_metadata_table(document: Document, metadata: dict[str, str]) -> None:
    """Render the report metadata block as a two-column table."""
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Mes/Ano", metadata.get("month", "")),
        ("Responsavel", metadata.get("responsible", "")),
        ("Equipe/Area", metadata.get("team", "")),
        ("Gerado em", dt.datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]
    for index, (label, value) in enumerate(rows):
        row = index // 2
        col = (index % 2) * 2
        if row >= len(table.rows):
            table.add_row()
        cell = table.cell(row, col)
        set_cell_text(cell, label, bold=True)
        if col + 1 >= len(table.columns):
            table.add_column(Inches(2.4))
        set_cell_text(table.cell(row, col + 1), value)


def add_cover_page(document: Document, title: str, metadata: dict[str, str]) -> None:
    """Create a simple cover section for the exported report."""
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run("Relatorio Interno")
    title_run.bold = True
    title_run.font.size = Pt(11)
    title_run.font.name = "Calibri"

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.add_run(title)
    heading_run.bold = True
    heading_run.font.size = Pt(20)
    heading_run.font.name = "Calibri"

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Documento preparado para distribuicao interna em PDF ou Word")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10.5)
    subtitle_run.font.name = "Calibri"

    document.add_paragraph()
    add_metadata_table(document, metadata)
    document.add_paragraph()
    add_horizontal_rule(document)


INLINE_PATTERN = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")


def add_inline_runs(paragraph, text: str) -> None:
    """Add text runs while honoring basic inline Markdown formatting."""
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            run.font.name = "Calibri"
        token = match.group(0)
        run_text = token
        run = paragraph.add_run()
        run.font.name = "Calibri"
        if token.startswith("`") and token.endswith("`"):
            run_text = token[1:-1]
            run.font.name = "Consolas"
        elif token.startswith("**") and token.endswith("**"):
            run_text = token[2:-2]
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run_text = token[1:-1]
            run.italic = True
        run.text = run_text
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        run.font.name = "Calibri"


def add_bullet_paragraph(document: Document, text: str, level: int) -> None:
    """Render a bullet item using a list style with a configurable nesting level."""
    style_name = "List Bullet"
    if level == 1:
        style_name = "List Bullet 2"
    elif level >= 2:
        style_name = "List Bullet 3"
    paragraph = document.add_paragraph(style=style_name)
    add_inline_runs(paragraph, text)


def add_code_block(document: Document, lines: list[str]) -> None:
    """Render a fenced code block with monospaced text."""
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["No Spacing"]
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)


def render_markdown_body(document: Document, text: str) -> None:
    """Render the Markdown body into DOCX paragraphs and headings."""
    in_code_block = False
    code_lines: list[str] = []

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                add_code_block(document, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(raw_line)
            continue

        if not stripped:
            document.add_paragraph()
            continue

        if stripped == "---":
            add_horizontal_rule(document)
            continue

        if line.startswith("### "):
            document.add_heading(stripped[4:], level=2)
            continue

        if line.startswith("## "):
            document.add_heading(stripped[3:], level=1)
            continue

        if line.startswith("# "):
            document.add_heading(stripped[2:], level=0)
            continue

        bullet_match = re.match(r"^(\s*)([-*])\s+(.+)$", line)
        if bullet_match:
            indent = len(bullet_match.group(1)) // 2
            add_bullet_paragraph(document, bullet_match.group(3), indent)
            continue

        paragraph = document.add_paragraph()
        add_inline_runs(paragraph, stripped)

    if code_lines:
        add_code_block(document, code_lines)


def add_summary_section(document: Document, toc: list[tuple[int, str, str]]) -> None:
    """Add a simple summary section from parsed headings."""
    document.add_heading("Sumario", level=1)
    for level, heading, _anchor in toc:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.2 * (level - 1))
        run = paragraph.add_run(heading)
        run.font.name = "Calibri"


def build_summary_document(title: str, metadata: dict[str, str], source_path: Path) -> str:
    """Create a short executive summary companion document."""
    return f"""Relatorio Executivo Interno
============================

Titulo: {title}
Mes/Ano: {metadata.get('month', '')}
Responsavel: {metadata.get('responsible', '')}
Equipe/Area: {metadata.get('team', '')}
Fonte: {source_path.name}
Gerado em: {dt.datetime.now().strftime('%Y-%m-%d %H:%M')}

Orientacao de uso:
- Abra o arquivo DOCX no Word.
- Use a opcao de salvar/exportar para PDF, se necessario.
- Compartilhe o PDF ou o DOCX exportado no drive interno.
"""


def build_docx_document(
    title: str,
    metadata: dict[str, str],
    toc: list[tuple[int, str, str]],
    body_source: str,
) -> Document:
    """Build the final DOCX document from the report contents."""
    document = Document()
    set_document_defaults(document)
    add_cover_page(document, title, metadata)
    add_summary_section(document, toc)
    document.add_paragraph()
    document.add_heading("Conteudo", level=1)
    render_markdown_body(document, body_source)
    return document


def main() -> int:
    """Run the export pipeline."""
    args = parse_args()
    input_path = Path(args.input).expanduser().resolve() if args.input else find_latest_report(DEFAULT_INPUT_DIR)
    output_dir = Path(args.output_dir).expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    source_text = input_path.read_text(encoding="utf-8")
    metadata = parse_report_metadata(source_text)
    toc = build_table_of_contents(source_text)
    report_title, body_source = extract_document_title_and_body(source_text, args.title)

    base_name = input_path.stem
    docx_path = output_dir / f"{base_name}.docx"
    summary_path = output_dir / f"{base_name}.summary.txt"

    document = build_docx_document(report_title, metadata, toc, body_source)
    document.save(docx_path)
    summary_path.write_text(
        build_summary_document(report_title, metadata, input_path),
        encoding="utf-8",
    )

    print(f"DOCX gerado: {docx_path}")
    print(f"Resumo gerado: {summary_path}")
    print("Se quiser distribuir, abra o DOCX no Word e exporte para PDF se necessario.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
